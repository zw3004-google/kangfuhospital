from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(r"D:\projects\kangfuhospital\康复医院系统部署及技术栈建议_信息科会议稿.docx")
BLUE = "1F4E78"
MID_BLUE = "D9EAF7"
LIGHT_BLUE = "EEF5FA"
LIGHT_GRAY = "F2F4F7"
GRAY = "666666"
RED = "9B1C1C"
GOLD = "7A5A00"
GREEN = "236B43"
FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_run(run, size=10.5, bold=False, color="222222"):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold_lead=None, color="222222", after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=color)
        set_run(p.add_run(text[len(bold_lead):]), color=color)
    else:
        set_run(p.add_run(text), color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths, header_fill=BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for i, value in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_fill)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=9.5, bold=True, color="FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], "F8FAFC")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=9.2)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    c = table.cell(0, 0)
    set_cell_shading(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(label + " "), bold=True, color=color)
    set_run(p.add_run(text))
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after in (("Heading 1", 15, 14, 7), ("Heading 2", 12.5, 10, 5), ("Heading 3", 11, 8, 4)):
    st = styles[name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("康复医院系统｜信息科会议确认稿"), size=8.5, color=GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer.add_run("内部讨论材料｜2026年8月"), size=8, color=GRAY)

# First page masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
set_run(p.add_run("部署服务器配置及技术栈建议"), size=23, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
set_run(p.add_run("康复医院欠费管理、预出院管理及系统管理平台"), size=12.5, color=GRAY)
add_table(doc, ["文档用途", "业务规模", "基础环境"], [["与医院信息科确认部署条件及技术路线", "单院内网；日导入约600～1000条；低至中并发", "虚拟机；宿主机无RAID；普通SAS机械盘"]], [2800, 3000, 3560])

doc.add_heading("一、结论摘要", level=1)
add_callout(doc, "建议结论", "采用“前后端分离 + Java模块化单体 + PostgreSQL主备”的技术路线。推荐4台生产虚拟机、1台测试虚拟机和1套独立备份存储。", LIGHT_BLUE, BLUE)
add_bullet(doc, "当前数据量不大，SAS机械盘在索引合理、报表预汇总的情况下可以满足基本性能。")
add_bullet(doc, "宿主机无RAID带来的首要问题是单盘故障和恢复风险，其次才是性能波动。")
add_bullet(doc, "主库与备库必须位于不同宿主机、不同物理SAS盘；否则多台虚拟机不具备实质容灾价值。")
add_bullet(doc, "数据库备份必须保存到第三套独立物理存储，虚拟机快照不能代替数据库备份。")

doc.add_heading("二、业务与容量假设", level=1)
add_table(doc, ["项目", "当前假设", "架构影响"], [
    ["业务模块", "欠费管理、预出院管理、系统管理", "适合模块化单体，不需要微服务"],
    ["数据导入", "每日约600～1000条，可能多次导入", "后台分批导入；保留导入批次和原始文件"],
    ["用户并发", "预计20～100人", "普通Web应用和关系型数据库足够"],
    ["报表", "排名、Top 3、趋势图、统计列表", "建立索引并采用定时汇总"],
    ["外部通信", "定时向企业微信推送并支持重发", "可靠任务表、幂等和失败重试"],
], [1700, 3500, 4160])

doc.add_heading("三、推荐部署方案", level=1)
add_table(doc, ["角色", "数量", "单台建议配置", "主要用途"], [
    ["应用服务器", "2台", "4 vCPU / 8～16GB / 150GB", "Vue静态资源、Java应用、定时与消息任务"],
    ["数据库主库", "1台", "8 vCPU / 32GB / 300～500GB", "PostgreSQL生产主库"],
    ["数据库备库", "1台", "8 vCPU / 32GB / 300～500GB", "PostgreSQL流复制热备"],
    ["测试服务器", "1台", "4 vCPU / 8～16GB / 150GB", "测试、升级与恢复验证"],
    ["独立备份存储", "1套", "500GB～1TB，可扩容", "数据库备份、WAL归档、原始导入文件"],
], [1800, 850, 3050, 3660])
add_text(doc, "生产虚拟机合计：4台；测试虚拟机：1台；另需1套不依赖生产宿主机磁盘的备份存储。", bold_lead="生产虚拟机合计：")

doc.add_heading("四、虚拟机放置与容灾前提", level=1)
add_table(doc, ["物理位置", "建议放置", "必须满足"], [
    ["宿主机A", "应用服务器1 + PostgreSQL主库", "主库所在SAS盘不承载高I/O的其他系统"],
    ["宿主机B", "应用服务器2 + PostgreSQL备库", "与主库不同宿主机、不同物理SAS盘"],
    ["独立存储/第三设备", "全量备份、WAL归档、导入原文件", "不能与主备库共用同一物理故障域"],
], [1900, 3450, 4010])
add_callout(doc, "关键风险", "如果所有虚拟机最终位于同一宿主机或同一块物理SAS盘，4台虚拟机仍然是单点系统。", "FDECEC", RED)

doc.add_heading("五、资源不足时的最低方案", level=1)
add_table(doc, ["服务器", "建议配置", "部署内容"], [
    ["生产虚拟机", "8 vCPU / 32GB / 500GB", "应用 + PostgreSQL主库"],
    ["备用虚拟机", "8 vCPU / 32GB / 500GB", "PostgreSQL备库；故障时人工启动应用"],
], [2200, 3000, 4160])
add_text(doc, "前提：两台虚拟机位于不同宿主机及不同物理盘，备份再复制到第三套独立设备。该方案能满足当前业务，但升级、维护和故障切换均不如推荐方案。")

doc.add_heading("六、SAS机械盘的预期影响与控制措施", level=1)
add_table(doc, ["影响", "具体表现", "控制措施"], [
    ["性能波动", "导入、报表、备份并发时页面可能明显变慢", "错峰备份；后台导入；报表定时汇总"],
    ["共享存储干扰", "其他虚拟机大量I/O时，本系统会突发卡顿", "数据库虚拟机尽量独占物理盘I/O资源"],
    ["单盘故障", "虚拟机不可用，主库可能整体丢失", "跨宿主机主备 + 第三方独立备份"],
    ["恢复时间", "需重建虚拟机、切换或恢复数据库", "制定切换手册并定期恢复演练"],
], [1650, 3700, 4010])
add_text(doc, "建议验收线：普通页面P95≤3秒；复杂报表P95≤8秒；1000条导入≤60秒；错误率＜1%；磁盘不应长期100%繁忙。", bold_lead="建议验收线：")

doc.add_heading("七、网络与安全要求", level=1)
add_table(doc, ["项目", "建议要求"], [
    ["院内访问", "统一通过HTTPS 443访问；用户终端不允许直连数据库"],
    ["服务器网络", "至少千兆内网；应用与数据库、主库与备库延迟尽量低于2ms"],
    ["企业微信", "应用服务器允许出站HTTPS 443、DNS和NTP；建议固定公网出口IP"],
    ["企微回调", "如需接收回调，应通过DMZ区HTTPS网关转发，不直接暴露内网应用"],
    ["访问控制", "数据库端口仅向应用服务器和备库开放；运维通过指定IP或堡垒机"],
], [2300, 7060])

doc.add_heading("八、推荐技术栈", level=1)
add_table(doc, ["层级", "推荐选型", "说明"], [
    ["总体架构", "前后端分离、模块化单体", "适配当前规模，降低部署和运维复杂度"],
    ["前端", "Vue 3 + TypeScript + Vite", "后台表格、表单及权限页面开发效率高"],
    ["UI与图表", "Element Plus + ECharts", "管理端组件及统计看板"],
    ["后端", "Java 21 + Spring Boot 3.5", "成熟稳定，适合医院长期维护"],
    ["权限", "Spring Security + RBAC", "菜单、按钮、数据范围及操作权限"],
    ["数据访问", "MyBatis / MyBatis-Plus", "便于复杂列表和统计SQL维护"],
    ["数据库", "PostgreSQL 17或医院支持版本", "事务、复杂统计和可靠性能力完整"],
    ["定时任务", "Spring Scheduler或Quartz", "定时汇总、企微推送及失败重试"],
    ["Excel", "EasyExcel", "分批导入、错误明细和批次追踪"],
    ["代理部署", "Nginx + Linux + Docker Compose/systemd", "无需Kubernetes"],
    ["缓存", "Redis（按需）", "登录状态和短期缓存；当前不是强制组件"],
], [1550, 3350, 4460])
add_callout(doc, "暂不建议", "一期不引入微服务、Kubernetes、Kafka或Elasticsearch，避免在机械盘与小规模业务下增加不必要的资源和运维负担。", "FFF7E1", GOLD)

doc.add_heading("九、关键实现约束", level=1)
add_table(doc, ["主题", "必须落实的设计"], [
    ["重复导入", "以院区 + 患者ID + 住院号/就诊号 + 住院次数作为业务唯一键；来源字段覆盖，人工标注字段保留"],
    ["导入审计", "保留导入批次、原始文件、导入人、时间、新增/更新/失败数量及错误明细"],
    ["并发编辑", "采用事务与乐观锁，避免多人处理时互相覆盖"],
    ["企微推送", "任务唯一编号、内容快照、状态、返回码、失败原因、发送次数、批量重发及幂等"],
    ["双应用实例", "定时任务采用数据库锁或Quartz集群，避免两台应用重复执行"],
    ["备份", "每日全量备份 + WAL持续归档；备份保留30～90天；定期恢复演练"],
], [2000, 7360])

doc.add_heading("十、会议需信息科确认的事项", level=1)
for item in [
    "是否能够提供至少两台不同宿主机，并确认主库、备库落在不同物理SAS盘？",
    "宿主机SAS盘的数量、转速、剩余容量，以及是否与其他高I/O系统共享？",
    "是否有独立NAS、备份服务器或其他第三方存储可保存数据库备份？",
    "是否允许Linux、Docker Compose、Nginx、PostgreSQL和Redis等软件？",
    "医院是否指定Java版本、数据库、中间件或存在信创适配要求？",
    "应用服务器能否访问企业微信公网接口，是否提供固定公网出口IP？",
    "是否需要接收企业微信回调；如需要，DMZ区和HTTPS证书如何提供？",
    "医院可接受的RPO、RTO分别是多少，是否要求自动切换？",
    "是否有等保等级、审计日志留存期限、数据备份期限等安全要求？",
    "测试环境、生产发布时间窗口、监控告警和日常运维责任由谁承担？",
]:
    add_bullet(doc, item)

add_callout(doc, "建议会议决策", "优先敲定物理故障域、备份位置和企微网络出口，再确认虚拟机规格。若这三项不明确，仅确认CPU、内存和虚拟机台数不足以形成可靠部署方案。", LIGHT_BLUE, BLUE)

doc.core_properties.title = "康复医院系统部署服务器配置及技术栈建议"
doc.core_properties.subject = "信息科会议确认稿"
doc.core_properties.author = "项目组"
doc.core_properties.keywords = "康复医院, 部署, 技术栈, PostgreSQL, 虚拟机"
doc.save(OUT)
print(OUT)
